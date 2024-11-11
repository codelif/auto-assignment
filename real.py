import openai
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import docx
from concurrent.futures import ThreadPoolExecutor

# Configuration
OPENAI_API_KEY = 'Api_key_here'  # Replace with your OpenAI API key
MODEL = 'gpt-4o-mini' 
TEMPDIR = tempfile.gettempdir()

# Set OpenAI API key
client = openai.Client(api_key=OPENAI_API_KEY)

def read_questions(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return [q.strip() for q in content.split('\n\n') if q.strip()]

def generate_solution(question):
    prompt = f"""   
    You are an expert C programmer. Provide a clear and concise solution to the following programming problem. Include the C code with proper formatting. Only output the code. DO NOT WRITE ANYTHING ELSE LIKE EXPLAINATION.

    Problem:
    {question}

    Solution:
    """
    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": "You are ChatGPT, an AI language model that provides detailed and accurate programming solutions."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.2,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Error generating solution for question: {question}\nError: {e}")
        return None

def extract_code(solution_text):
    import re
    code_pattern = re.compile(r'```c(.*?)```', re.DOTALL)
    match = code_pattern.search(solution_text)
    if match:
        return match.group(1).strip()
    else:
        code_pattern = re.compile(r'```(.*?)```', re.DOTALL)
        match = code_pattern.search(solution_text)
        if match:
            return match.group(1).strip()
    return None

def compile_and_run_c_code(c_code):
    """
    Compiles and runs the provided C code.
    Returns the output or error messages, simulating input if necessary.
    """
    simulated_input = "42\n"  # Example input; modify as needed for your programs.
    with tempfile.NamedTemporaryFile(delete=False, suffix='.c', dir=TEMPDIR) as c_file:
        c_file_name = c_file.name
        c_file.write(c_code.encode('utf-8'))

    exe_file = c_file_name.replace('.c', '.exe') if os.name == 'nt' else c_file_name.replace('.c', '')

    # Compile the C code
    compile_cmd = ['gcc', c_file_name, '-o', exe_file]
    try:
        compile_process = subprocess.run(compile_cmd, capture_output=True, text=True, timeout=10)
        if compile_process.returncode != 0:
            os.unlink(c_file_name)
            return {"status": "error", "message": f"Compilation failed:\n{compile_process.stderr}"}
    except subprocess.TimeoutExpired:
        os.unlink(c_file_name)
        return {"status": "error", "message": "Compilation timed out during compilation."}

    # Run the executable with simulated input
    run_cmd = [exe_file]
    try:
        run_process = subprocess.run(run_cmd, capture_output=True, text=True, input=simulated_input, timeout=3)
        output = run_process.stdout
        if run_process.stderr:
            output += f"\nRuntime Errors:\n{run_process.stderr}"
        return {"status": "success", "message": output.strip()}
    except subprocess.TimeoutExpired:
        return {"status": "error", "message": "Execution timed out during runtime."}
    except Exception as e:
        return {"status": "error", "message": f"Error during execution: {e}"}
    finally:
        os.unlink(c_file_name)
        if os.path.exists(exe_file):
            os.unlink(exe_file)
            

def create_docx(qa_list, output_path):
    document = Document()
    document.add_heading('C Programming Solutions', 0)

    for idx, qa in enumerate(qa_list, 1):
        question = qa['question']
        solution = qa['solution']
        output = qa['output']

        # Add Question
        document.add_heading(f'Question {idx}', level=1)
        document.add_paragraph(question)

        # Add Solution with Code
        document.add_heading('Solution', level=2)
        code = extract_code(solution)
        if code:
            explanation = solution.replace(f'```c\n{code}\n```', '').strip()
            document.add_paragraph(explanation)
            document.add_paragraph('C Code:')
            code_paragraph = document.add_paragraph()
            code_paragraph.style.font.name = 'Consolas'
            run = code_paragraph.add_run(code)
            run.font.size = Pt(10)
        else:
            document.add_paragraph(solution)

        # Add Output with Black Background and White Text
        document.add_heading('Output', level=2)
        output_paragraph = document.add_paragraph()
        output_paragraph.style.font.name = 'Consolas'
        run = output_paragraph.add_run(output if isinstance(output, str) else output['message'])
        run.font.size = Pt(10)
        run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)  # Set font color to white

        # Set black background for the output text
        pPr = output_paragraph._element.get_or_add_pPr()  # Access paragraph properties
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "000000")  # Set background to black
        pPr.append(shading_elm)

        document.add_page_break()

    document.save(output_path)
    print(f"Document saved to {output_path}")



def process_question(question):
    solution = generate_solution(question)
    if not solution:
        return {'question': question, 'solution': "Failed to generate solution.", 'output': "N/A"}
    
    code = extract_code(solution)
    output = compile_and_run_c_code(code) if code else "No executable code found in the solution."
    return {'question': question, 'solution': solution, 'output': output}

def main():
    if len(sys.argv) != 2:
        print("Usage: python generate_solutions.py questions.txt")
        sys.exit(1)

    questions_file = sys.argv[1]
    if not os.path.exists(questions_file):
        print(f"File not found: {questions_file}")
        sys.exit(1)

    questions = read_questions(questions_file)
    print(f"Found {len(questions)} questions.")

    with ThreadPoolExecutor() as executor:
        qa_list = list(executor.map(process_question, questions))

    output_docx = 'C_Programming_Solutions.docx'
    create_docx(qa_list, output_docx)

if __name__ == '__main__':
    main()
